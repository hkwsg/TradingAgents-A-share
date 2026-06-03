"""将茅台分析报告 Markdown 转换为排版精美的 Word 文档"""
import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SRC = Path(os.environ["USERPROFILE"]) / "Desktop" / "TradingAgent报告_600519_2026-05-07" / "完整分析报告.md"
DST = Path(os.environ["USERPROFILE"]) / "Desktop" / "贵州茅台_600519_交易分析报告.docx"

doc = Document()

# -- 页面设置 --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(4)

# 预先定义标题样式
for i in range(1, 4):
    h_style = doc.styles[f'Heading {i}']
    h_font = h_style.font
    h_font.name = '微软雅黑'
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if i == 1:
        h_font.size = Pt(20)
        h_font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    elif i == 2:
        h_font.size = Pt(15)
        h_font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)
    else:
        h_font.size = Pt(12)
        h_font.color.rgb = RGBColor(0x3D, 0x3D, 0x3D)

def add_colored_table(doc, headers, rows):
    """创建带交替行色的美观表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="333333"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # 数据行
    for r, row_data in enumerate(rows):
        row = table.rows[r + 1]
        for c, text in enumerate(row_data):
            cell = row.cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(text))
            run.font.size = Pt(9)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            if r % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
                cell._element.get_or_add_tcPr().append(shading)

    doc.add_paragraph()  # 表后间距
    return table

def parse_inline(doc, p, text):
    """解析行内格式：粗体、斜体、代码"""
    # 先处理 `code`
    parts = re.split(r'(`[^`]+`)', text)
    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            # 处理 **bold** 和 *italic*
            sub_parts = re.split(r'(\*\*[^*]+\*\*)', part)
            for sp in sub_parts:
                if sp.startswith('**') and sp.endswith('**'):
                    run = p.add_run(sp[2:-2])
                    run.bold = True
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                else:
                    # 处理 *italic*
                    italic_parts = re.split(r'(\*[^*]+\*)', sp)
                    for ip in italic_parts:
                        if ip.startswith('*') and ip.endswith('*') and not ip.startswith('**'):
                            run = p.add_run(ip[1:-1])
                            run.italic = True
                            run.font.name = '微软雅黑'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                        else:
                            run = p.add_run(ip)
                            run.font.name = '微软雅黑'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

lines = SRC.read_text(encoding='utf-8').split('\n')

i = 0
in_table = False
table_rows = []
table_headers = []

while i < len(lines):
    line = lines[i]

    # 跳过空行
    if not line.strip():
        i += 1
        continue

    # 水平分隔线
    if line.strip().startswith('---'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>')
        pPr.append(pBdr)
        i += 1
        continue

    # 标题
    if line.startswith('# '):
        doc.add_heading(line[2:].strip(), level=1)
        i += 1
        continue
    if line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
        i += 1
        continue
    if line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
        i += 1
        continue
    if line.startswith('#### '):
        doc.add_heading(line[5:].strip(), level=3)
        i += 1
        continue

    # 表格检测
    if line.strip().startswith('|') and line.strip().endswith('|'):
        if not in_table:
            in_table = True
            table_rows = []
            # 表头行
            headers = [c.strip() for c in line.split('|')[1:-1]]
            table_headers = headers
            i += 1
            # 跳过分隔行 (|:---|:---|)
            if i < len(lines) and '---' in lines[i]:
                i += 1
            continue
        else:
            # 数据行
            cells = [c.strip() for c in line.split('|')[1:-1]]
            table_rows.append(cells)
            i += 1
            # 检查下一行是否还是表格
            if i >= len(lines) or not (lines[i].strip().startswith('|') and lines[i].strip().endswith('|')):
                add_colored_table(doc, table_headers, table_rows)
                in_table = False
                table_rows = []
                table_headers = []
            continue

    # 列表项
    if re.match(r'^\s*[\-\*\d+\.]\s', line):
        text = re.sub(r'^\s*[\-\*\d+\.]\s+', '', line)
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        parse_inline(doc, p, text)
        i += 1
        continue

    # 引用块（"数据源访问失败"等）
    if line.strip().startswith('> '):
        text = line.strip()[2:]
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.italic = True
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.size = Pt(9)
        i += 1
        continue

    # 普通段落
    p = doc.add_paragraph()
    parse_inline(doc, p, line)
    i += 1

# 添加页脚说明
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 本报告由 TradingAgents 多智能体交易分析框架自动生成，仅供参考，不构成投资建议 —')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

doc.save(str(DST))
print(f"Word doc saved: {DST}")
