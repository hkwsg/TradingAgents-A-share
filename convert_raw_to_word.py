"""从 markdown 报告提取各章节，按项目原生结构生成 Word"""
import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SRC = Path(os.environ["USERPROFILE"]) / "Desktop" / "TradingAgent报告_600519_2026-05-07" / "完整分析报告.md"
DST = Path(os.environ["USERPROFILE"]) / "Desktop" / "贵州茅台_600519_原始分析输出.docx"

text = SRC.read_text(encoding='utf-8')

# ---- 拆分章节 ----
sections = {}
current_section = "前言"
current_content = []

for line in text.split('\n'):
    if line.startswith('# 一、分析师团队报告'):
        sections[current_section] = '\n'.join(current_content)
        current_section = "analysts"
        current_content = []
    elif line.startswith('# 二、研究团队决策'):
        sections[current_section] = '\n'.join(current_content)
        current_section = "research"
        current_content = []
    elif line.startswith('# 三、交易员执行计划'):
        sections[current_section] = '\n'.join(current_content)
        current_section = "trader"
        current_content = []
    elif line.startswith('# 四、组合经理最终决策'):
        sections[current_section] = '\n'.join(current_content)
        current_section = "portfolio"
        current_content = []
    else:
        current_content.append(line)
sections[current_section] = '\n'.join(current_content)

# ---- 从 analysts 中提取子报告 ----
def extract_subsections(content):
    """从内容中提取子章节"""
    result = []
    lines = content.split('\n')
    current = {"title": "", "body": []}
    for line in lines:
        if line.startswith('## ') and ('技术' in line or '情绪' in line or '新闻' in line or '基本面' in line or '社交' in line):
            if current["title"]:
                current["body"] = '\n'.join(current["body"])
                result.append(current)
            current = {"title": line[3:], "body": []}
        else:
            current["body"].append(line)
    if current["title"]:
        current["body"] = '\n'.join(current["body"])
        result.append(current)
    return result

analyst_subs = extract_subsections(sections.get("analysts", ""))

# ---- 构建 Word ----
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.35

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hf = hs.font
    hf.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    colors = {1: 0x1A1A1A, 2: 0x2B5EAA, 3: 0x3D3D3D}
    sizes = {1: 18, 2: 14, 3: 12}
    hf.size = Pt(sizes[i])
    hf.color.rgb = RGBColor(*[(colors[i] >> 16) & 0xFF, (colors[i] >> 8) & 0xFF, colors[i] & 0xFF])

def add_para(text, bold=False, size=10.5, indent=0, color=None, align=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    return p

def add_block(text, is_quote=False):
    """添加文本块，只处理粗体和引用"""
    p = doc.add_paragraph()
    if is_quote:
        p.paragraph_format.left_indent = Cm(1)
    for line in text.split('\n'):
        if not line.strip():
            continue
        # 粗体处理
        parts = re.split(r'(\*\*[^*]+\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part if part else ' ')
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(9.5) if is_quote else Pt(10.5)
            if is_quote:
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p.add_run('\n')
    return p

# ======== 封面 ========
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n')

add_para('TradingAgents 多智能体交易分析', bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x1A, 0x1A, 0x1A))
add_para('原始分析输出', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x2B, 0x5E, 0xAA))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('—' * 40)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

info_lines = [
    ('标的', '600519.SH — 贵州茅台'),
    ('分析日期', '2026-05-07'),
    ('模型', 'deepseek-chat (快速) / deepseek-v4-pro (深度)'),
    ('输出语言', 'Chinese'),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = p.add_run(f'{label}：')
    run_label.font.size = Pt(10)
    run_label.bold = True
    run_label.font.name = '微软雅黑'
    run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run_value = p.add_run(value)
    run_value.font.size = Pt(10)
    run_value.font.name = '微软雅黑'
    run_value._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ======== 目录概览 ========
doc.add_heading('输出结构说明', level=1)
add_para('本项目模拟真实交易公司的团队架构，分析流程分四个阶段，每阶段输出独立报告：')

toc_items = [
    '1_Analyst Team — 分析师团队：市场技术、社会情绪、新闻、基本面 四份独立报告',
    '2_Research Team — 研究团队：多空辩论 + 研究经理裁决',
    '3_Trading Team — 交易员：基于A股T+1制度的执行计划',
    '4_Portfolio Management — 组合经理：综合风控后的最终决策 + 信号提取',
]
for item in toc_items:
    p = doc.add_paragraph(style='List Number')
    p.clear()
    run = p.add_run(item)
    run.font.size = Pt(10)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ======== 阶段一：分析师团队 ========
doc.add_heading('阶段一：分析师团队报告', level=1)

analyst_subtitles = {
    '市场技术分析': 'Market Analyst — 市场技术分析',
    '社会情绪分析': 'Social Analyst — 社会情绪分析',
    '新闻分析': 'News Analyst — 新闻分析',
    '基本面分析': 'Fundamentals Analyst — 基本面分析',
}

for sub in analyst_subs:
    doc.add_heading(sub["title"], level=2)
    # 提取关键摘要（前200字）
    body = sub["body"]
    # 简化处理：按段落添加
    paras = body.split('\n\n')
    for para in paras:
        para = para.strip()
        if not para:
            continue
        if para.startswith('|') and para.endswith('|'):
            # 简单表格处理
            rows = [r for r in para.split('\n') if r.strip().startswith('|') and r.strip().endswith('|')]
            if len(rows) >= 2:
                hdr = [c.strip() for c in rows[0].split('|')[1:-1]]
                data = [[c.strip() for c in r.split('|')[1:-1]] for r in rows[2:]]
                if hdr and data:
                    table = doc.add_table(rows=1+len(data), cols=len(hdr))
                    table.style = 'Table Grid'
                    for i, h in enumerate(hdr):
                        cell = table.rows[0].cells[i]
                        cell.text = h
                        for p2 in cell.paragraphs:
                            for r2 in p2.runs:
                                r2.bold = True
                                r2.font.size = Pt(8)
                    for ri, rd in enumerate(data):
                        for ci, cv in enumerate(rd):
                            cell = table.rows[ri+1].cells[ci]
                            cell.text = cv
                            for p2 in cell.paragraphs:
                                for r2 in p2.runs:
                                    r2.font.size = Pt(8)
                    doc.add_paragraph()
        elif para.startswith('#'):
            doc.add_heading(para.strip('#').strip(), level=3)
        else:
            add_block(para)

doc.add_page_break()

# ======== 阶段二：研究团队 ========
doc.add_heading('阶段二：研究团队辩论与决策', level=1)

research = sections.get("research", "")
research_parts = research.split('### ')
for part in research_parts:
    part = part.strip()
    if not part:
        continue
    lines = part.split('\n', 1)
    title = lines[0].strip()
    body = lines[1].strip() if len(lines) > 1 else ''
    if title in ('多头观点', '空头观点', '研究经理裁决'):
        doc.add_heading(title, level=2)
        add_block(body)

doc.add_page_break()

# ======== 阶段三：交易员 ========
doc.add_heading('阶段三：交易员执行计划', level=1)
trader = sections.get("trader", "")
trader = re.sub(r'^## .*\n', '', trader).strip()
add_block(trader)

doc.add_page_break()

# ======== 阶段四：组合经理 ========
doc.add_heading('阶段四：组合经理最终决策', level=1)
portfolio = sections.get("portfolio", "")
portfolio = re.sub(r'^## .*\n', '', portfolio).strip()
add_block(portfolio)

# 信号提取
signal_match = re.search(r'\*\*提取的交易信号\*\*: `(\w+)`', text)
if signal_match:
    signal = signal_match.group(1)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'最终交易信号：{signal}')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 页脚
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 本报告由 TradingAgents 多智能体交易分析框架自动生成，仅供参考，不构成投资建议 —')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

doc.save(str(DST))
print(f"Raw output doc saved: {DST}")
